#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
📎 Упрощенный обработчик вложений для уже отфильтрованных файлов
Работает только с файлами, которые прошли раннюю фильтрацию в fetcher'е
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, List, Optional


class AttachmentProcessor:
    """📎 Упрощенный процессор для отфильтрованных вложений"""
    
    def __init__(self):
        print("📎 Инициализация упрощенного обработчика вложений")
        print("   ✅ Работает только с предварительно отфильтрованными файлами")
        print("   🇷🇺 OCR с поддержкой русского + английского языков")

    def process_email_attachments(self, email: Dict, email_loader) -> Dict:
        """📧 Обработка ТОЛЬКО отфильтрованных вложений"""
        
        attachments = email.get('attachments', [])
        if not attachments:
            return {
                'attachments_processed': 0,
                'attachments_text': [],
                'processing_errors': []
            }
        
        # Фильтруем только фактически скачанные файлы
        downloaded_attachments = [
            att for att in attachments 
            if att.get('status') == 'saved' and att.get('file_path')
        ]
        
        print(f"   📎 К обработке: {len(downloaded_attachments)} из {len(attachments)} вложений")
        
        # Показываем что было исключено при скачивании
        excluded_attachments = [att for att in attachments if att.get('status') == 'excluded']
        unsupported_attachments = [att for att in attachments if att.get('status') == 'unsupported']
        
        if excluded_attachments:
            print(f"   🚫 Исключено при скачивании: {len(excluded_attachments)}")
            for att in excluded_attachments:
                print(f"      • {att['original_filename']} - {att['exclusion_reason']}")
        
        if unsupported_attachments:
            print(f"   ⚠️ Неподдерживаемых при скачивании: {len(unsupported_attachments)}")
            for att in unsupported_attachments:
                print(f"      • {att['original_filename']} - {att['exclusion_reason']}")
        
        # Обрабатываем только скачанные файлы
        processed_attachments = []
        processing_errors = []
        
        for i, attachment in enumerate(downloaded_attachments, 1):
            try:
                result = self.process_single_attachment(email, attachment, email_loader)
                if result:
                    processed_attachments.append(result)
                    print(f"      ✅ {i}/{len(downloaded_attachments)}: {attachment.get('original_filename', 'N/A')}")
                else:
                    print(f"      ⚠️ {i}/{len(downloaded_attachments)}: Не удалось обработать {attachment.get('original_filename', 'N/A')}")
            except Exception as e:
                error_msg = f"Ошибка обработки {attachment.get('original_filename', 'N/A')}: {e}"
                processing_errors.append(error_msg)
                print(f"      ❌ {i}/{len(downloaded_attachments)}: {error_msg}")
        
        return {
            'attachments_processed': len(processed_attachments),
            'attachments_text': processed_attachments,
            'processing_errors': processing_errors
        }

    def process_single_attachment(self, email: Dict, attachment: Dict, email_loader) -> Optional[Dict]:
        """📎 Обработка одного вложения"""
        
        file_path = email_loader.get_attachment_file_path(email, attachment)
        if not file_path or not file_path.exists():
            return None
        
        file_type = attachment.get('file_type', '')
        original_filename = attachment.get('original_filename', '')
        file_size = attachment.get('file_size', 0)
        
        # Определяем метод обработки по типу файла
        if file_type == 'application/pdf':
            processing_method = 'pdf_with_ocr'
        elif file_type.startswith('image/'):
            processing_method = 'ocr'
        elif file_type == 'text/plain':
            processing_method = 'direct'
        elif 'excel' in file_type or 'sheet' in file_type:
            processing_method = 'office_excel'
        elif 'word' in file_type:
            processing_method = 'office_word'
        else:
            processing_method = 'unsupported'
        
        # Извлекаем текст в зависимости от типа
        if processing_method == 'direct':
            extracted_text = self._extract_text_direct(file_path)
        elif processing_method == 'pdf_with_ocr':
            extracted_text = self._extract_text_from_pdf_with_ocr(file_path, original_filename)
        elif processing_method == 'ocr':
            extracted_text = self._extract_text_with_ocr(file_path, original_filename)
        elif processing_method == 'office_excel':
            extracted_text = self._extract_from_excel(file_path, original_filename)
        elif processing_method == 'office_word':
            extracted_text = self._extract_from_word(file_path, original_filename)
        else:
            extracted_text = f"[НЕПОДДЕРЖИВАЕМЫЙ ФОРМАТ: {file_type}]"
        
        return {
            'filename': original_filename,
            'file_type': file_type,
            'method': processing_method,
            'extracted_text': extracted_text,
            'char_count': len(extracted_text),
            'file_size': file_size
        }

    def _extract_text_direct(self, file_path: Path) -> str:
        """📄 Прямое извлечение текста из текстовых файлов"""
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return content.strip()
        except Exception as e:
            return f"[ОШИБКА ЧТЕНИЯ ФАЙЛА: {e}]"

    def _extract_text_from_pdf_with_ocr(self, file_path: Path, filename: str) -> str:
        """📄 Извлечение текста из PDF с поддержкой OCR для сканов"""
        
        try:
            # Сначала пробуем извлечь обычный текст из PDF
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_parts = []
                ocr_needed_pages = []
                
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text = page.get_text()
                    
                    if text.strip() and len(text.strip()) > 50:
                        # В PDF есть текстовый слой
                        text_parts.append(f"[СТРАНИЦА {page_num + 1} - ТЕКСТОВЫЙ СЛОЙ]\n{text}")
                    else:
                        # Страница скорее всего отсканированная - нужен OCR
                        ocr_needed_pages.append(page_num)
                
                # Если есть страницы для OCR
                if ocr_needed_pages:
                    print(f"        🔍 PDF содержит сканированные страницы: {len(ocr_needed_pages)} из {doc.page_count}")
                    
                    # Применяем OCR к сканированным страницам
                    for page_num in ocr_needed_pages:
                        ocr_result = self._extract_pdf_page_with_ocr(doc, page_num, filename)
                        if ocr_result:
                            text_parts.append(f"[СТРАНИЦА {page_num + 1} - OCR ИЗВЛЕЧЕНИЕ]\n{ocr_result}")
                
                doc.close()
                
                if text_parts:
                    return '\n\n'.join(text_parts)
                else:
                    return f"[PDF БЕЗ ТЕКСТОВОГО СОДЕРЖИМОГО: {filename}]"
                    
            except ImportError:
                return f"[PDF ФАЙЛ: {filename} - ТРЕБУЕТСЯ PyMuPDF ДЛЯ ОБРАБОТКИ]\nУстановите: pip install PyMuPDF"
            
        except Exception as e:
            return f"[ОШИБКА ОБРАБОТКИ PDF: {e}]"

    def _extract_pdf_page_with_ocr(self, doc, page_num: int, filename: str) -> str:
        """🔍 OCR обработка страницы PDF"""
        
        try:
            # Импортируем библиотеки для OCR
            try:
                import pytesseract
                from PIL import Image
                import io
            except ImportError:
                return f"[OCR БИБЛИОТЕКИ НЕ УСТАНОВЛЕНЫ - Страница {page_num + 1}]"
            
            # Получаем страницу как изображение
            page = doc[page_num]
            # Увеличиваем разрешение для лучшего OCR (300 DPI)
            mat = page.get_pixmap(matrix=page.get_matrix(zoom=2.0))
            img_data = mat.tobytes("png")
            image = Image.open(io.BytesIO(img_data))
            
            # OCR с поддержкой русского и английского
            extracted_text = pytesseract.image_to_string(
                image,
                lang='rus+eng',  # 🇷🇺 Русский + английский
                config='--psm 6'
            )
            
            if extracted_text.strip():
                return self._clean_ocr_text(extracted_text.strip())
            else:
                return f"[OCR НЕ СМОГ ИЗВЛЕЧЬ ТЕКСТ СО СТРАНИЦЫ {page_num + 1}]"
                
        except Exception as e:
            return f"[ОШИБКА OCR СТРАНИЦЫ {page_num + 1}: {e}]"

    def _extract_text_with_ocr(self, file_path: Path, filename: str) -> str:
        """🖼️ Реальная OCR обработка изображений с поддержкой русского"""
        
        try:
            # Импортируем библиотеки для OCR
            try:
                import pytesseract
                from PIL import Image, ImageEnhance
            except ImportError:
                return f"""[OCR БИБЛИОТЕКИ НЕ УСТАНОВЛЕНЫ]
Файл: {filename}
Для OCR обработки установите: pip install pytesseract pillow"""
            
            print(f"        🔍 Запуск OCR для {filename}...")
            
            # Открываем изображение
            image = Image.open(file_path)
            
            # Предобработка изображения для лучшего OCR
            # Увеличиваем контрастность
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            
            # Увеличиваем четкость
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2.0)
            
            # OCR с поддержкой русского и английского языков
            extracted_text = pytesseract.image_to_string(
                image, 
                lang='rus+eng',  # 🇷🇺 Русский + английский
                config='--psm 6'
            )
            
            if extracted_text.strip():
                print(f"        ✅ OCR извлечено {len(extracted_text)} символов")
                
                # Постобработка текста
                cleaned_text = self._clean_ocr_text(extracted_text)
                
                return f"""[OCR ИЗВЛЕЧЕННЫЙ ТЕКСТ]
Файл: {filename}
Языки: Русский + Английский
Символов извлечено: {len(cleaned_text)}

{cleaned_text}"""
            else:
                return f"""[OCR НЕ СМОГ ИЗВЛЕЧЬ ТЕКСТ]
Файл: {filename}
Возможные причины: низкое качество изображения, неразборчивый текст"""
                
        except Exception as e:
            return f"""[ОШИБКА OCR ОБРАБОТКИ]
Файл: {filename}
Ошибка: {e}"""

    def _clean_ocr_text(self, text: str) -> str:
        """🧹 Очистка текста после OCR"""
        
        # Убираем лишние пробелы и переносы
        text = re.sub(r'\s+', ' ', text)
        
        # Убираем одиночные символы на отдельных строках
        text = re.sub(r'\n\s*[^\w\s]\s*\n', '\n', text)
        
        # Объединяем разорванные слова (если OCR разбил слово переносом)
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        
        # Убираем повторяющиеся знаки препинания
        text = re.sub(r'([.!?]){2,}', r'\1', text)
        
        return text.strip()

    def _extract_from_excel(self, file_path: Path, filename: str) -> str:
        """📊 Извлечение текста из Excel файлов"""
        
        try:
            import openpyxl
            import xlrd
            
            text_parts = []
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.xlsx':
                # Обрабатываем XLSX через openpyxl
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"[ЛИСТ: {sheet_name}]")
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell in row:
                            if cell is not None:
                                row_text.append(str(cell))
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
            elif file_ext == '.xls':
                # Обрабатываем XLS через xlrd
                workbook = xlrd.open_workbook(file_path)
                
                for sheet_idx in range(workbook.nsheets):
                    sheet = workbook.sheet_by_index(sheet_idx)
                    text_parts.append(f"[ЛИСТ: {sheet.name}]")
                    
                    for row_idx in range(sheet.nrows):
                        row_text = []
                        for col_idx in range(sheet.ncols):
                            cell = sheet.cell(row_idx, col_idx)
                            if cell.value:
                                row_text.append(str(cell.value))
                        if row_text:
                            text_parts.append(" | ".join(row_text))
            
            if text_parts:
                return f"""[EXCEL ФАЙЛ ОБРАБОТАН]
Файл: {filename}
Листов: {len([p for p in text_parts if p.startswith('[ЛИСТ:')])}

{chr(10).join(text_parts)}"""
            else:
                return f"[EXCEL ФАЙЛ ПУСТ: {filename}]"
                
        except ImportError:
            return f"""[EXCEL БИБЛИОТЕКИ НЕ УСТАНОВЛЕНЫ]
Файл: {filename}
Установите: pip install openpyxl xlrd"""
        except Exception as e:
            return f"[ОШИБКА ЧТЕНИЯ EXCEL: {e}]"

    def _extract_from_word(self, file_path: Path, filename: str) -> str:
        """📝 Извлечение текста из Word файлов"""
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_parts = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                text_parts.append("[ТАБЛИЦА]")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if text_parts:
                return f"""[WORD ДОКУМЕНТ ОБРАБОТАН]
Файл: {filename}
Параграфов: {len([p for p in text_parts if not p.startswith('[')])}
Таблиц: {len([p for p in text_parts if p == '[ТАБЛИЦА]'])}

{chr(10).join(text_parts)}"""
            else:
                return f"[WORD ДОКУМЕНТ ПУСТ: {filename}]"
                
        except ImportError:
            return f"""[WORD БИБЛИОТЕКА НЕ УСТАНОВЛЕНА]
Файл: {filename}
Установите: pip install python-docx"""
        except Exception as e:
            return f"[ОШИБКА ЧТЕНИЯ WORD: {e}]"

    def combine_email_with_attachments(self, email: Dict, attachments_result: Dict) -> str:
        """🔗 Объединение текста письма с содержимым вложений"""
        
        combined_parts = []
        
        # Основной текст письма
        email_body = email.get('body', '').strip()
        if email_body:
            combined_parts.append("=== ТЕКСТ ПИСЬМА ===")
            combined_parts.append(email_body)
        
        # Информация об исключенных/неподдерживаемых вложениях (из JSON письма)
        attachments = email.get('attachments', [])
        excluded_attachments = [att for att in attachments if att.get('status') == 'excluded']
        unsupported_attachments = [att for att in attachments if att.get('status') == 'unsupported']
        
        if excluded_attachments or unsupported_attachments:
            combined_parts.append("\n=== ИНФОРМАЦИЯ ОБ ИСКЛЮЧЕННЫХ ВЛОЖЕНИЯХ ===")
            for att in excluded_attachments:
                combined_parts.append(f"🚫 {att['original_filename']} - {att['exclusion_reason']}")
            for att in unsupported_attachments:
                combined_parts.append(f"⚠️ {att['original_filename']} - {att['exclusion_reason']}")
        
        # Текст из обработанных вложений
        attachments_text = attachments_result.get('attachments_text', [])
        if attachments_text:
            combined_parts.append("\n=== СОДЕРЖИМОЕ ВЛОЖЕНИЙ ===")
            
            for i, attachment in enumerate(attachments_text, 1):
                filename = attachment.get('filename', f'Вложение {i}')
                method = attachment.get('method', 'unknown')
                extracted_text = attachment.get('extracted_text', '')
                
                combined_parts.append(f"\n--- ВЛОЖЕНИЕ {i}: {filename} (метод: {method}) ---")
                combined_parts.append(extracted_text)
        
        # Информация об ошибках
        processing_errors = attachments_result.get('processing_errors', [])
        if processing_errors:
            combined_parts.append("\n=== ОШИБКИ ОБРАБОТКИ ВЛОЖЕНИЙ ===")
            for error in processing_errors:
                combined_parts.append(f"❌ {error}")
        
        return '\n'.join(combined_parts)


def main():
    """🧪 Тестирование упрощенного обработчика вложений"""
    
    print("📎 ТЕСТИРОВАНИЕ УПРОЩЕННОГО ОБРАБОТЧИКА ВЛОЖЕНИЙ")
    print("="*70)
    
    # Импортируем загрузчик писем
    import sys
    sys.path.append(str(Path(__file__).parent))
    from email_loader import ProcessedEmailLoader
    
    # Загружаем письма для тестирования
    loader = ProcessedEmailLoader()
    available_dates = loader.get_available_date_folders()
    
    if not available_dates:
        print("❌ Нет обработанных писем для тестирования")
        return
    
    # Берем последнюю доступную дату
    latest_date = available_dates[-1]
    emails = loader.load_emails_by_date(latest_date)
    
    # Фильтруем письма с вложениями
    emails_with_attachments = loader.get_emails_with_attachments(emails)
    
    if not emails_with_attachments:
        print("❌ Нет писем с вложениями для тестирования")
        return
    
    print(f"📎 Найдено писем с вложениями: {len(emails_with_attachments)}")
    
    # Создаем процессор
    processor = AttachmentProcessor()
    
    # Статистика по всем письмам
    total_processed = 0
    total_extracted = 0
    
    # Тестируем на первых 3 письмах с вложениями
    for i, email in enumerate(emails_with_attachments[:3], 1):
        print(f"\n{'='*70}")
        print(f"📧 ПИСЬМО {i}: {email.get('subject', 'N/A')[:50]}...")
        print(f"📨 От: {email.get('from', 'N/A')[:40]}...")
        
        # Показываем статистику вложений из JSON
        attachments_stats = email.get('attachments_stats', {})
        print(f"📊 Статистика из JSON: {attachments_stats}")
        
        # Обрабатываем вложения
        attachments_result = processor.process_email_attachments(email, loader)
        
        total_processed += attachments_result['attachments_processed']
        
        # Показываем результат
        print(f"✅ Успешно обработано: {attachments_result['attachments_processed']}")
        
        # Показываем детали каждого вложения
        for j, att_result in enumerate(attachments_result.get('attachments_text', []), 1):
            print(f"   📎 Вложение {j}: {att_result['filename']}")
            print(f"      Метод: {att_result['method']}")
            print(f"      Символов: {att_result['char_count']}")
            
            if att_result['char_count'] > 0:
                total_extracted += 1
                # Показываем превью извлеченного текста
                preview = att_result['extracted_text'][:200]
                print(f"      Превью: {preview}...")
        
        # Объединяем с текстом письма
        combined_text = processor.combine_email_with_attachments(email, attachments_result)
        
        print(f"📝 Общий объем текста: {len(combined_text)} символов")
    
    # Итоговая статистика
    print(f"\n{'='*70}")
    print(f"📊 ИТОГОВАЯ СТАТИСТИКА ОБРАБОТКИ")
    print(f"{'='*70}")
    print(f"📎 Всего вложений обработано: {total_processed}")
    print(f"✅ Успешно извлечен текст: {total_extracted}")
    
    if total_processed > 0:
        success_rate = (total_extracted / total_processed) * 100
        print(f"📈 Успешность извлечения: {success_rate:.1f}%")


if __name__ == '__main__':
    main()
